"""
Utilitaires pour exporter les résultats WHOIS en Word et Excel
"""
from io import BytesIO
from datetime import datetime
from django.http import HttpResponse


def export_whois_to_word(domain, ip_address, country_from_tld, domain_info, raw_whois):
    """
    Exporte les résultats WHOIS en format Word (.docx)
    """
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    doc = Document()
    
    # Titre du document
    title = doc.add_heading('Rapport WHOIS', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_format = title.runs[0]
    title_format.font.color.rgb = RGBColor(0, 102, 255)
    
    # Informations de génération
    info = doc.add_paragraph()
    info.add_run(f"Domaine: ").bold = True
    info.add_run(f"{domain}\n")
    info.add_run(f"Pays (TLD): ").bold = True
    info.add_run(f"{country_from_tld}\n")
    info.add_run(f"Date du rapport: ").bold = True
    info.add_run(f"{datetime.now().strftime('%d/%m/%Y %H:%M:%S')}\n")
    info.add_run(f"Outil: ").bold = True
    info.add_run("Protegio WHOIS Tool\n")
    
    # Adresse IP
    doc.add_heading('Résolution IP', level=1)
    ip_para = doc.add_paragraph()
    ip_para.add_run("Adresse IP: ").bold = True
    ip_para.add_run(f"{ip_address if ip_address else 'Non résolue'}")
    
    # Informations du domaine
    if domain_info:
        doc.add_heading('Informations du Domaine', level=1)
        
        info_table = doc.add_table(rows=1, cols=2)
        info_table.style = 'Light Grid Accent 1'
        
        # En-têtes
        header_cells = info_table.rows[0].cells
        header_cells[0].text = "Propriété"
        header_cells[1].text = "Valeur"
        
        # Données
        data_fields = [
            ("Nom du Domaine", domain_info.get('domain_name', 'N/A')),
            ("Registraire", domain_info.get('registrar', 'N/A')),
            ("Date de Création", str(domain_info.get('creation_date', 'N/A'))),
            ("Date d'Expiration", str(domain_info.get('expiration_date', 'N/A'))),
            ("Dernière Mise à Jour", str(domain_info.get('last_updated', 'N/A'))),
            ("Pays", domain_info.get('country', 'N/A')),
            ("Organisation", domain_info.get('org', 'N/A')),
        ]
        
        for field, value in data_fields:
            row = info_table.add_row()
            row.cells[0].text = field
            row.cells[1].text = str(value)
        
        # Serveurs DNS
        doc.add_heading('Serveurs DNS', level=2)
        nameservers = domain_info.get('name_servers', [])
        if nameservers:
            for ns in nameservers:
                doc.add_paragraph(str(ns), style='List Bullet')
        else:
            doc.add_paragraph("Aucun serveur DNS disponible")
        
        # Statut
        doc.add_heading('Statut du Domaine', level=2)
        status = domain_info.get('status', [])
        if status:
            for st in status:
                doc.add_paragraph(str(st), style='List Bullet')
        else:
            doc.add_paragraph("Aucun statut disponible")
    
    # WHOIS brut
    if raw_whois:
        doc.add_heading('Données WHOIS Brutes', level=1)
        raw_para = doc.add_paragraph(raw_whois)
        raw_para.paragraph_format.style = 'Normal'
    
    # Ajouter un pied de page
    doc.add_page_break()
    footer = doc.add_paragraph()
    footer.add_run("Généré par Protegio WHOIS Tool").italic = True
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Ajouter un style au document
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    # Préparer la réponse
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="WHOIS_{domain}_{datetime.now().strftime("%d%m%Y_%H%M%S")}.docx"'
    
    doc.save(response)
    return response


def export_whois_to_excel(domain, ip_address, country_from_tld, domain_info, raw_whois):
    """
    Exporte les résultats WHOIS en format Excel (.xlsx)
    """
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    
    wb = openpyxl.Workbook()
    
    # Première feuille: Résumé
    ws = wb.active
    ws.title = "Résumé"
    
    # Styles
    header_fill = PatternFill(start_color="0066FF", end_color="0066FF", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF")
    border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )
    
    # En-têtes
    ws.column_dimensions['A'].width = 25
    ws.column_dimensions['B'].width = 50
    
    # Titre
    ws['A1'] = "RAPPORT WHOIS"
    ws['A1'].font = Font(bold=True, size=14, color="0066FF")
    ws.merge_cells('A1:B1')
    
    # Informations générales
    row = 3
    ws[f'A{row}'] = "Domaine"
    ws[f'B{row}'] = domain
    row += 1
    
    ws[f'A{row}'] = "Pays (TLD)"
    ws[f'B{row}'] = country_from_tld
    row += 1
    
    ws[f'A{row}'] = "Date du rapport"
    ws[f'B{row}'] = datetime.now().strftime('%d/%m/%Y %H:%M:%S')
    row += 1
    
    ws[f'A{row}'] = "Adresse IP"
    ws[f'B{row}'] = ip_address if ip_address else 'Non résolue'
    row += 2
    
    # Informations du domaine
    if domain_info:
        ws[f'A{row}'] = "INFORMATIONS DU DOMAINE"
        ws[f'A{row}'].font = Font(bold=True, size=11, color="0066FF")
        ws.merge_cells(f'A{row}:B{row}')
        row += 1
        
        data_fields = [
            ("Nom du Domaine", domain_info.get('domain_name', 'N/A')),
            ("Registraire", domain_info.get('registrar', 'N/A')),
            ("Date de Création", str(domain_info.get('creation_date', 'N/A'))),
            ("Date d'Expiration", str(domain_info.get('expiration_date', 'N/A'))),
            ("Dernière Mise à Jour", str(domain_info.get('last_updated', 'N/A'))),
            ("Pays", domain_info.get('country', 'N/A')),
            ("Organisation", domain_info.get('org', 'N/A')),
        ]
        
        for field, value in data_fields:
            ws[f'A{row}'] = field
            ws[f'B{row}'] = str(value)
            ws[f'A{row}'].font = Font(bold=True)
            row += 1
        
        # Serveurs DNS
        row += 1
        ws[f'A{row}'] = "SERVEURS DNS"
        ws[f'A{row}'].font = Font(bold=True, size=11, color="0066FF")
        ws.merge_cells(f'A{row}:B{row}')
        row += 1
        
        nameservers = domain_info.get('name_servers', [])
        if nameservers:
            for ns in nameservers:
                ws[f'A{row}'] = str(ns)
                row += 1
        else:
            ws[f'A{row}'] = "Aucun serveur DNS disponible"
            row += 1
        
        # Statut du domaine
        row += 1
        ws[f'A{row}'] = "STATUT DU DOMAINE"
        ws[f'A{row}'].font = Font(bold=True, size=11, color="0066FF")
        ws.merge_cells(f'A{row}:B{row}')
        row += 1
        
        status = domain_info.get('status', [])
        if status:
            for st in status:
                ws[f'A{row}'] = str(st)
                row += 1
        else:
            ws[f'A{row}'] = "Aucun statut disponible"
            row += 1
    
    # Feuille 2: Données brutes WHOIS
    if raw_whois:
        ws2 = wb.create_sheet("Données Brutes")
        ws2['A1'] = "DONNÉES WHOIS BRUTES"
        ws2['A1'].font = Font(bold=True, size=12, color="0066FF")
        
        ws2.column_dimensions['A'].width = 150
        
        # Diviser le texte brut en lignes
        lines = raw_whois.split('\n')
        for idx, line in enumerate(lines, start=3):
            ws2[f'A{idx}'] = line
            ws2.row_dimensions[idx].height = 15
    
    # Préparer la réponse
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = f'attachment; filename="WHOIS_{domain}_{datetime.now().strftime("%d%m%Y_%H%M%S")}.xlsx"'
    
    wb.save(response)
    return response


def export_all_whois_to_word(results):
    """
    Exporte TOUS les résultats WHOIS en format Word avec une page par résultat
    """
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    doc = Document()
    
    # Page de couverture
    title = doc.add_heading('Rapport WHOIS Complet', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_format = title.runs[0]
    title_format.font.color.rgb = RGBColor(0, 102, 255)
    
    info = doc.add_paragraph()
    info.add_run(f"Date du rapport: ").bold = True
    info.add_run(f"{datetime.now().strftime('%d/%m/%Y %H:%M:%S')}\n")
    info.add_run(f"Nombre de domaines: ").bold = True
    info.add_run(f"{results.count()}\n")
    info.add_run(f"Outil: ").bold = True
    info.add_run("Protegio WHOIS Tool")
    
    # Chaque résultat sur une nouvelle page
    for idx, result in enumerate(results, 1):
        if idx > 1:
            doc.add_page_break()
        
        # Titre du domaine
        domain_title = doc.add_heading(f'Domaine {idx}: {result.domain}', level=1)
        domain_title_format = domain_title.runs[0]
        domain_title_format.font.color.rgb = RGBColor(0, 102, 255)
        
        # Métadonnées
        meta = doc.add_paragraph()
        meta.add_run(f"Dernière mise à jour: ").bold = True
        meta.add_run(f"{result.updated_at.strftime('%d/%m/%Y %H:%M:%S')}\n")
        
        # Adresse IP
        doc.add_heading('Résolution IP', level=2)
        ip_para = doc.add_paragraph()
        ip_para.add_run("Adresse IP: ").bold = True
        ip_para.add_run(f"{result.ip_address if result.ip_address else 'Non résolue'}")
        
        # Informations du domaine
        if result.domain_info:
            doc.add_heading('Informations du Domaine', level=2)
            
            info_table = doc.add_table(rows=1, cols=2)
            info_table.style = 'Light Grid Accent 1'
            
            header_cells = info_table.rows[0].cells
            header_cells[0].text = "Propriété"
            header_cells[1].text = "Valeur"
            
            data_fields = [
                ("Nom du Domaine", result.domain_info.get('domain_name', 'N/A')),
                ("Registraire", result.domain_info.get('registrar', 'N/A')),
                ("Date de Création", str(result.domain_info.get('creation_date', 'N/A'))),
                ("Date d'Expiration", str(result.domain_info.get('expiration_date', 'N/A'))),
                ("Dernière Mise à Jour", str(result.domain_info.get('last_updated', 'N/A'))),
                ("Pays", result.domain_info.get('country', 'N/A')),
                ("Organisation", result.domain_info.get('org', 'N/A')),
            ]
            
            for field, value in data_fields:
                row = info_table.add_row()
                row.cells[0].text = field
                row.cells[1].text = str(value)
            
            # Serveurs DNS
            doc.add_heading('Serveurs DNS', level=3)
            nameservers = result.domain_info.get('name_servers', [])
            if nameservers:
                for ns in nameservers:
                    doc.add_paragraph(str(ns), style='List Bullet')
            else:
                doc.add_paragraph("Aucun serveur DNS disponible")
            
            # Statut
            doc.add_heading('Statut du Domaine', level=3)
            status = result.domain_info.get('status', [])
            if status:
                for st in status:
                    doc.add_paragraph(str(st), style='List Bullet')
            else:
                doc.add_paragraph("Aucun statut disponible")
        
        # WHOIS brut
        if result.raw_whois:
            doc.add_heading('Données WHOIS Brutes', level=2)
            raw_para = doc.add_paragraph(result.raw_whois)
            raw_para.paragraph_format.style = 'Normal'
    
    # Ajouter un pied de page final
    doc.add_page_break()
    footer = doc.add_paragraph()
    footer.add_run("Généré par Protegio WHOIS Tool").italic = True
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="WHOIS_Complet_{datetime.now().strftime("%d%m%Y_%H%M%S")}.docx"'
    
    doc.save(response)
    return response


def export_all_whois_to_excel(results):
    """
    Exporte TOUS les résultats WHOIS en format Excel avec plusieurs feuilles
    """
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Résumé"
    
    # Feuille de résumé
    ws['A1'] = "RAPPORT WHOIS COMPLET"
    ws['A1'].font = Font(bold=True, size=14, color="0066FF")
    ws.merge_cells('A1:D1')
    
    ws.column_dimensions['A'].width = 30
    ws.column_dimensions['B'].width = 20
    ws.column_dimensions['C'].width = 20
    ws.column_dimensions['D'].width = 30
    
    row = 3
    headers = ["Domaine", "Adresse IP", "Registraire", "Dernière Mise à Jour"]
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=row, column=col)
        cell.value = header
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill(start_color="0066FF", end_color="0066FF", fill_type="solid")
    
    row = 4
    for result in results:
        ws.cell(row=row, column=1).value = result.domain
        ws.cell(row=row, column=2).value = result.ip_address or "N/A"
        ws.cell(row=row, column=3).value = result.domain_info.get('registrar', 'N/A')
        ws.cell(row=row, column=4).value = result.updated_at.strftime('%d/%m/%Y %H:%M:%S')
        row += 1
    
    # Créer une feuille par domaine
    for result in results:
        ws_domain = wb.create_sheet(result.domain[:31])  # Limite à 31 caractères
        
        ws_domain.column_dimensions['A'].width = 25
        ws_domain.column_dimensions['B'].width = 50
        
        # En-tête
        ws_domain['A1'] = f"WHOIS - {result.domain}"
        ws_domain['A1'].font = Font(bold=True, size=12, color="0066FF")
        ws_domain.merge_cells('A1:B1')
        
        row = 3
        
        # Adresse IP
        ws_domain[f'A{row}'] = "Adresse IP"
        ws_domain[f'B{row}'] = result.ip_address or "Non résolue"
        ws_domain[f'A{row}'].font = Font(bold=True)
        row += 1
        
        ws_domain[f'A{row}'] = "Date de mise à jour"
        ws_domain[f'B{row}'] = result.updated_at.strftime('%d/%m/%Y %H:%M:%S')
        ws_domain[f'A{row}'].font = Font(bold=True)
        row += 2
        
        # Informations du domaine
        if result.domain_info:
            ws_domain[f'A{row}'] = "INFORMATIONS DU DOMAINE"
            ws_domain[f'A{row}'].font = Font(bold=True, size=11, color="0066FF")
            ws_domain.merge_cells(f'A{row}:B{row}')
            row += 1
            
            data_fields = [
                ("Nom du Domaine", result.domain_info.get('domain_name', 'N/A')),
                ("Registraire", result.domain_info.get('registrar', 'N/A')),
                ("Date de Création", str(result.domain_info.get('creation_date', 'N/A'))),
                ("Date d'Expiration", str(result.domain_info.get('expiration_date', 'N/A'))),
                ("Dernière Mise à Jour", str(result.domain_info.get('last_updated', 'N/A'))),
                ("Pays", result.domain_info.get('country', 'N/A')),
                ("Organisation", result.domain_info.get('org', 'N/A')),
            ]
            
            for field, value in data_fields:
                ws_domain[f'A{row}'] = field
                ws_domain[f'B{row}'] = str(value)
                ws_domain[f'A{row}'].font = Font(bold=True)
                row += 1
            
            # Serveurs DNS
            row += 1
            ws_domain[f'A{row}'] = "SERVEURS DNS"
            ws_domain[f'A{row}'].font = Font(bold=True, size=11, color="0066FF")
            ws_domain.merge_cells(f'A{row}:B{row}')
            row += 1
            
            nameservers = result.domain_info.get('name_servers', [])
            if nameservers:
                for ns in nameservers:
                    ws_domain[f'A{row}'] = str(ns)
                    row += 1
            else:
                ws_domain[f'A{row}'] = "Aucun serveur DNS"
                row += 1
            
            # Statut
            row += 1
            ws_domain[f'A{row}'] = "STATUT"
            ws_domain[f'A{row}'].font = Font(bold=True, size=11, color="0066FF")
            ws_domain.merge_cells(f'A{row}:B{row}')
            row += 1
            
            status = result.domain_info.get('status', [])
            if status:
                for st in status:
                    ws_domain[f'A{row}'] = str(st)
                    row += 1
            else:
                ws_domain[f'A{row}'] = "Aucun statut"
    
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = f'attachment; filename="WHOIS_Complet_{datetime.now().strftime("%d%m%Y_%H%M%S")}.xlsx"'
    
    wb.save(response)
    return response
